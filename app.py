from datetime import datetime
from io import BytesIO
from flask import Flask, render_template, render_template_string, request, send_file
from docx import Document
import os
import smtplib
from email.message import EmailMessage
from email.mime.text import MIMEText

app = Flask(__name__)

CAMINHO_MODELO = os.path.join(os.path.dirname(__file__), 'Procuração Pessoa Física.docx')

def enviar_email_com_anexo(destino, assunto, corpo, arquivo_bytes, nome_arquivo):

    msg = EmailMessage()
    msg['Subject'] = assunto
    msg['From'] = "envioemail.gen@gmail.com"
    msg['To'] = destino
    msg.set_content(corpo)
    msg.add_attachment(arquivo_bytes,
                       maintype='application',
                       subtype='vnd.openxmlformats-officedocument.wordprocessingml.document',
                       filename=nome_arquivo)
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login("envioemail.gen@gmail.com", "djmo cvwn bsba kjom")
        smtp.send_message(msg)

@app.route('/', methods=['GET', 'POST'])
def form():
    if request.method == 'POST':
        # Formata a data de nascimento para o formato desejado
        data_nascimento_raw = request.form.get("data_nascimento", "")
        try:
            # converte a string para datetime
            data_nascimento_dt = datetime.strptime(data_nascimento_raw, "%Y-%m-%d")
            # agora pode usar strftime
            data_nascimento_formatada = data_nascimento_dt.strftime("%d/%m/%Y")
        except (ValueError, TypeError):
            # Se não der pra converter, mantém como está ou define vazio
            data_nascimento_formatada = data_nascimento_raw

        # Formata CPF
        cpf_raw = request.form.get("cpf", "")
        if len(cpf_raw) == 11:
            cpf_formatado = f"{cpf_raw[:3]}.{cpf_raw[3:6]}.{cpf_raw[6:9]}-{cpf_raw[9:]}"

        else:
            cpf_formatado = cpf_raw

        rg_raw = request.form.get("rg", "")
        if len(rg_raw) == 9:
            rg_formatado = f"{rg_raw[:2]}.{rg_raw[2:5]}.{rg_raw[5:8]}-{rg_raw[8]}"
        else:
            rg_formatado = rg_raw

        nome_completo =  request.form.get("nome_completo", "")
        nome_formatado = " ".join([
            palavra[0].upper() + palavra[1:].lower() if palavra else ''
            for palavra in nome_completo.split()
        ])

        profissao = request.form.get("profissao", "")
        profissao_formatada = " ".join([
            palavra[0].upper() + palavra[1:].lower() if palavra else ''
            for palavra in profissao.split()
        ])

        cidade_nascimento = request.form.get("cidade_nascimento", "")
        cidade_formatada = " ".join([
            palavra[0].upper() + palavra[1:].lower() if palavra else ''
            for palavra in cidade_nascimento.split()
        ])

        nome_pai = request.form.get("nome_pai", "")
        nome_pai_formatado = " ".join([
            palavra[0].upper() + palavra[1:].lower() if palavra else ''
            for palavra in nome_pai.split()
        ])

        nome_mae = request.form.get("nome_pai", "")
        nome_mae_formatado = " ".join([
            palavra[0].upper() + palavra[1:].lower() if palavra else ''
            for palavra in nome_mae.split()
        ])
            
        # 1) Captura dos dados do formulário
        tipo_rg = request.form.get("rg_tipo", "")
        if tipo_rg == "rg_antigo":
                rg_valor = ", e da Carteira de Identidade Registro Geral n  " + rg_formatado + "/"
                orgao_emissor_valor = request.form.get("orgao_emissor", "") + "/"
                estado_emissor_valor = request.form.get("estado_emissor", "")
        else:
                rg_valor = ""
                orgao_emissor_valor = ""
                estado_emissor_valor = ""
        dados = {
            "<<NOME>>":                             nome_formatado,
            "<<ESTADO CIVIL>>":                     request.form.get("estado_civil", ""),
            "<<NACIONALIDADE>>":                    request.form.get("nacionalidade", ""),
            "<<CPF>>":                              cpf_formatado,
            "<<DATA NASCIMENTO>>":                  data_nascimento_formatada,
            "<<PROFISSAO>>":                        profissao_formatada,
            "<<CIDADE NASCIMENTO>>":                cidade_formatada,
            "<<ESTADO NASCIMENTO>>":                request.form.get("estado_nascimento", ""),
            "<<RG>>":                               rg_valor,
            "<<ORGAO EMISSOR>>":                    orgao_emissor_valor,
            "<<ESTADO EMISSOR>>":                   estado_emissor_valor,
            "<<NOME PAI>>":                         nome_pai_formatado,
            "<<NOME MAE>>":                         nome_mae_formatado,
            "<<CIDADE DOMICILIO>>":                 request.form.get("cidade", ""),
            "<<ESTADO DOMICILIO>>":                 request.form.get("estado", ""),
            "<<LOGRADOURO DOMICILIO>>":             request.form.get("logradouro", ""),
            "<<RUA DOMICILIO>>":                    request.form.get("rua", ""),
            "<<NUMERO DOMICILIO>>":                 request.form.get("numero", ""),
            "<<BAIRRO DOMICILIO>>":                 request.form.get("bairro", ""),
            "<<CEP>>":                              request.form.get("cep", ""),
            "<<GENERO>>": "a"                    if request.form.get("genero", "") == "feminino" else "o",
        }

        # 2) Captura data e hora de preenchimento
        data_preenchimento = datetime.now().strftime("%d/%m/%Y")
        # Você pode injetar esse valor em um placeholder, por exemplo:
        dados["<<DATA PREENCHIMENTO>>"] = data_preenchimento

        if not os.path.isfile(CAMINHO_MODELO):
            return f"Arquivo modelo não encontrado em: {CAMINHO_MODELO}", 500

        # Abre e preenche o documento
        doc = Document(CAMINHO_MODELO)
        def substituir(p):
            for run in p.runs:
                for ph, val in dados.items():
                    if ph in run.text:
                        run.text = run.text.replace(ph, val)

        for p in doc.paragraphs:
            substituir(p)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        substituir(p)

        # Gera o nome e salva no servidor
        nome_pessoa = request.form.get("nome_completo", "documento").strip().replace(" ", "_")
        nome_saida = f"{nome_pessoa}_Procuração_PF.docx"

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        arquivo_bytes = file_stream.read()

        # Envia email automático pra vocês
        enviar_email_com_anexo(
            destino='candidoemartins.adv@outlook.com',
            assunto=f"Nova procuração gerada - {nome_pessoa}",
            corpo='Segue em anexo o documento gerado pelo formulário.',
            arquivo_bytes=arquivo_bytes,
            nome_arquivo=nome_saida
        )
      
    return render_template('form.html')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)